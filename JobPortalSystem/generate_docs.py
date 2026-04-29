from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ── helpers ───────────────────────────────────────────────────────────────────
def h(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def para(text="", bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    return p

def mono(text):
    p = doc.add_paragraph()
    p.style = doc.styles['No Spacing']
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(8.5)
    return p

def gap(): doc.add_paragraph()

def shade_row(row, hex_color="1F4E79"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                r.bold = True

def table(rows_data, col_widths=None):
    t = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    t.style = 'Table Grid'
    for i, row_data in enumerate(rows_data):
        row = t.rows[i]
        for j, val in enumerate(row_data):
            c = row.cells[j]
            c.text = str(val)
            c.paragraphs[0].runs[0].font.size = Pt(9)
            if i == 0: c.paragraphs[0].runs[0].bold = True
        if i == 0: shade_row(row)
    return t

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
gap(); gap()
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("Job Portal & Resume Builder")
r.bold=True; r.font.size=Pt(26); r.font.color.rgb=RGBColor(0x1F,0x4E,0x79)

tp2 = doc.add_paragraph(); tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run("System Design Document")
r2.bold=True; r2.font.size=Pt(18); r2.font.color.rgb=RGBColor(0x2E,0x74,0xB5)

gap()
tp3 = doc.add_paragraph(); tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tp3.add_run("Activity Diagrams  ·  Sequence Diagrams  ·  Class Diagrams  ·  UML  ·  ER Diagrams")
r3.italic=True; r3.font.size=Pt(12)
gap(); gap()

# ══════════════════════════════════════════════════════════════════════════════
# 1. ARCHITECTURE OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h("1. Architecture Overview")
para("Seven .NET 10 microservices, each owning its own SQL Server database.", italic=True, size=10)
gap()
table([
    ("Service","Port","Database","Responsibility"),
    ("ApiGateway","5000","—","Ocelot routing, JWT RS256 validation, rate limiting"),
    ("IdentityService","5001","IdentityDb","Auth, JWT RS256, refresh tokens, account lockout"),
    ("JobCatalogService","5002","JobDb","Job listings, Elasticsearch search, companies"),
    ("ApplicationService","5003","ApplicationDb","Applications, state machine, email notifications"),
    ("ResumeService","5004","ResumeDb","Resume CRUD, PDF generation, recruiter unlock"),
    ("PaymentService","5005","PaymentDb","Razorpay wallet, points deduction, webhooks"),
    ("AdminService","5006","AdminDb","User/job moderation, audit logs, revenue reports"),
])
gap()
h("Communication Channels", 2)
for l in [
    "• HTTP (sync)   : ResumeService → PaymentService  (points deduction for resume view)",
    "• HTTP (sync)   : AdminService → IdentityDb / JobDb / PaymentDb  (direct cross-DB reads)",
    "• RabbitMQ      : ApplicationService publishes ApplicationStatusChanged  (fanout, durable)",
    "• Elasticsearch : JobCatalogService indexes / searches job documents",
    "• Redis         : IdentityService stores JTI blocklist for logout revocation",
]:
    para(l, size=10)

# ══════════════════════════════════════════════════════════════════════════════
# 2. ER DIAGRAMS  (one per service database)
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h("2. Entity-Relationship (ER) Diagrams")
para("Each service owns its own database. Cross-service references are by GUID only (no FK constraints across DBs).", italic=True, size=10)

# ── 2.1 IdentityDb ────────────────────────────────────────────────────────────
h("2.1  IdentityDb", 2)
for l in """
+---------------------------+          +-----------------------------+
|          Users            |          |        RefreshTokens        |
+---------------------------+          +-----------------------------+
| PK  Id            GUID    |1        *| PK  Id            GUID      |
|     Email         VARCHAR |----------| FK  UserId        GUID      |
|     PasswordHash  VARCHAR |          |     TokenHash     VARCHAR   |
|     Role          VARCHAR |          |     ExpiresAt     DATETIME  |
|     DisplayName   VARCHAR |          |     IsRevoked     BIT       |
|     EmailVerified BIT     |          |     CreatedAt     DATETIME  |
|     EmailVerificationToken|          +-----------------------------+
|     PasswordResetToken    |
|     FailedLoginAttempts INT|
|     LockoutEnd    DATETIME|
|     IsDeleted     BIT     |
|     CreatedAt     DATETIME|
|     UpdatedAt     DATETIME|
+---------------------------+

Relationships:
  Users (1) ──────────── (*) RefreshTokens   [UserId FK, CASCADE DELETE]
""".strip().splitlines():
    mono(l)
gap()

# ── 2.2 JobDb ─────────────────────────────────────────────────────────────────
h("2.2  JobDb", 2)
for l in """
+---------------------------+       +-----------------------------+
|        Companies          |       |           Jobs              |
+---------------------------+       +-----------------------------+
| PK  Id          GUID      |1     *| PK  Id              GUID    |
|     Name        VARCHAR   |-------| FK  CompanyId       GUID    |
|     Description TEXT      |       |     PostedByRecruiter GUID  |
|     Website     VARCHAR   |       |     Title           VARCHAR |
|     LogoUrl     VARCHAR   |       |     Description     TEXT    |
|     Industry    VARCHAR   |       |     Location        VARCHAR |
|     Location    VARCHAR   |       |     JobType         VARCHAR |
|     CreatedAt   DATETIME  |       |     SalaryMin       DECIMAL |
+---------------------------+       |     SalaryMax       DECIMAL |
         |                          |     ModerationStatus VARCHAR|
         |1                         |     CreatedAt       DATETIME|
         |                          |     UpdatedAt       DATETIME|
         |*                         +-----------------------------+
+---------------------------+
|      CompanyReviews       |
+---------------------------+
| PK  Id          GUID      |
| FK  CompanyId   GUID      |
|     ReviewerId  GUID      |
|     Rating      INT       |
|     Comment     TEXT      |
|     IsApproved  BIT       |
|     CreatedAt   DATETIME  |
+---------------------------+

Relationships:
  Companies (1) ──── (*) Jobs            [CompanyId FK]
  Companies (1) ──── (*) CompanyReviews  [CompanyId FK]
""".strip().splitlines():
    mono(l)
gap()

# ── 2.3 ApplicationDb ─────────────────────────────────────────────────────────
h("2.3  ApplicationDb", 2)
for l in """
+-------------------------------+
|         Applications          |
+-------------------------------+
| PK  Id              GUID      |
|     JobSeekerId     GUID      |  ← references IdentityDb.Users (by GUID)
|     JobId           GUID      |  ← references JobDb.Jobs (by GUID)
|     JobSeekerEmail  VARCHAR   |
|     Status          VARCHAR   |  Submitted|Reviewed|Shortlisted|Rejected|Withdrawn
|     CoverLetter     TEXT      |
|     AppliedAt       DATETIME  |
|     UpdatedAt       DATETIME  |
+-------------------------------+

Unique Index: (JobSeekerId, JobId)  — prevents duplicate applications
Index:        (JobSeekerId, Status)
Index:        (JobId, Status)
""".strip().splitlines():
    mono(l)
gap()

# ── 2.4 ResumeDb ──────────────────────────────────────────────────────────────
h("2.4  ResumeDb", 2)
for l in """
+-------------------------+    +-------------------------+
|         Resumes         |    |      ResumeEducation    |
+-------------------------+    +-------------------------+
| PK Id         GUID      |1  *| PK Id         GUID      |
|    OwnerId    GUID       |---| FK ResumeId   GUID      |
|    Title      VARCHAR    |   |    Institution VARCHAR  |
|    Summary    TEXT       |   |    Degree      VARCHAR  |
|    TemplateId VARCHAR    |   |    FieldOfStudy VARCHAR |
|    Certifications TEXT   |   |    StartDate   DATETIME |
|    CreatedAt  DATETIME   |   |    EndDate     DATETIME |
|    UpdatedAt  DATETIME   |   +-------------------------+
+-------------------------+
         |1                    +-------------------------+
         |--------------------*|     ResumeExperience    |
         |                     +-------------------------+
         |                     | PK Id         GUID      |
         |                     | FK ResumeId   GUID      |
         |                     |    Company    VARCHAR   |
         |                     |    JobTitle   VARCHAR   |
         |                     |    Description TEXT     |
         |                     |    StartDate  DATETIME  |
         |                     |    EndDate    DATETIME  |
         |                     |    IsCurrentRole BIT    |
         |                     +-------------------------+
         |1
         |--------------------*+-------------------------+
         |                     |       ResumeSkill       |
         |                     +-------------------------+
         |                     | PK Id       GUID        |
         |                     | FK ResumeId GUID        |
         |                     |    Name     VARCHAR     |
         |                     |    Level    VARCHAR     |
         |                     +-------------------------+
         |1
         |--------------------*+-------------------------+
                               |      ResumeProject      |
                               +-------------------------+
                               | PK Id       GUID        |
                               | FK ResumeId GUID        |
                               |    Name     VARCHAR     |
                               |    Description TEXT     |
                               |    Url      VARCHAR     |
                               +-------------------------+

+---------------------------+    +---------------------------+
|   ResumeUnlockRequests    |    |      UnlockedResumes      |
+---------------------------+    +---------------------------+
| PK Id          GUID       |    | PK Id          GUID       |
|    RecruiterId GUID       |    |    RecruiterId GUID       |
|    ResumeId    GUID       |    |    ResumeId    GUID       |
|    Status      VARCHAR    |    |    UnlockedAt  DATETIME   |
|    CreatedAt   DATETIME   |    +---------------------------+
|    UpdatedAt   DATETIME   |    Unique Index: (RecruiterId, ResumeId)
+---------------------------+

All child tables (Education/Experience/Skill/Project) CASCADE DELETE on ResumeId.
""".strip().splitlines():
    mono(l)
gap()

# ── 2.5 PaymentDb ─────────────────────────────────────────────────────────────
h("2.5  PaymentDb", 2)
for l in """
+-----------------------------+       +-----------------------------+
|       RecruiterWallets      |       |         Transactions        |
+-----------------------------+       +-----------------------------+
| PK  Id            GUID      |1     *| PK  Id              GUID    |
|     RecruiterId   GUID      |-------| FK  WalletId        GUID    |
|     PointsBalance INT       |       |     Type            VARCHAR | Credit|Debit
|     CreatedAt     DATETIME  |       |     Points          INT     |
|     UpdatedAt     DATETIME  |       |     Description     TEXT    |
|     RowVersion    TIMESTAMP |       |     IdempotencyKey  VARCHAR | UNIQUE
+-----------------------------+       |     RazorpayPaymentId VARCHAR|
                                      |     AmountPaid      DECIMAL |
                                      |     Currency        VARCHAR |
                                      |     CreatedAt       DATETIME|
                                      +-----------------------------+

+-----------------------------+       +-----------------------------+
|     PointsDeductionRules    |       |    UnlockResumeSagaState    |
+-----------------------------+       +-----------------------------+
| PK  Id       GUID           |       | PK  CorrelationId  GUID     |
|     Action   VARCHAR        |       |     CurrentState   VARCHAR  |
|     Points   INT            |       |     RecruiterId    GUID     |
|     IsActive BIT            |       |     ResumeId       GUID     |
+-----------------------------+       |     PointsToDeduct INT      |
                                      |     CreatedAt      DATETIME |
                                      |     UpdatedAt      DATETIME |
                                      +-----------------------------+

Relationships:
  RecruiterWallets (1) ──── (*) Transactions  [WalletId FK]
  Unique Index: RecruiterWallets.RecruiterId
  Unique Index: Transactions.IdempotencyKey  (WHERE NOT NULL)
""".strip().splitlines():
    mono(l)
gap()

# ── 2.6 AdminDb ───────────────────────────────────────────────────────────────
h("2.6  AdminDb", 2)
for l in """
+-----------------------------+
|          AuditLogs          |
+-----------------------------+
| PK  Id          GUID        |
|     AdminId     GUID        |  ← references IdentityDb.Users (by GUID)
|     Action      VARCHAR     |  e.g. "ApproveJob", "FlagJob", "ChangeRole"
|     TargetType  VARCHAR     |  e.g. "Job", "User"
|     TargetId    GUID        |
|     Details     TEXT        |
|     CreatedAt   DATETIME    |
+-----------------------------+

Index: (AdminId, CreatedAt)
Note: AdminDb is append-only. No updates or deletes on AuditLogs.
""".strip().splitlines():
    mono(l)
gap()

# ══════════════════════════════════════════════════════════════════════════════
# 3. CLASS DIAGRAMS  (one per service)
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h("3. Class Diagrams")
para("UML class notation: + public  # protected  - private  <<interface>> stereotype", italic=True, size=10)

# ── 3.1 IdentityService ───────────────────────────────────────────────────────
h("3.1  IdentityService", 2)
for l in """
┌─────────────────────────────────────────┐
│              <<entity>>                 │
│                 User                    │
├─────────────────────────────────────────┤
│ + Id                  : Guid            │
│ + Email               : string          │
│ + PasswordHash        : string          │
│ + Role                : string          │
│ + DisplayName         : string?         │
│ + EmailVerified       : bool            │
│ + EmailVerificationToken : string?      │
│ + EmailVerificationTokenExpiry: DateTime?│
│ + PasswordResetToken  : string?         │
│ + PasswordResetTokenExpiry : DateTime?  │
│ + FailedLoginAttempts : int             │
│ + LockoutEnd          : DateTime?       │
│ + IsDeleted           : bool            │
│ + CreatedAt           : DateTime        │
│ + UpdatedAt           : DateTime        │
│ + RefreshTokens       : ICollection<>   │
└─────────────────────────────────────────┘
                    1 │
                      │ has many
                    * ▼
┌─────────────────────────────────────────┐
│              <<entity>>                 │
│             RefreshToken                │
├─────────────────────────────────────────┤
│ + Id         : Guid                     │
│ + UserId     : Guid  (FK)               │
│ + User       : User                     │
│ + TokenHash  : string                   │
│ + ExpiresAt  : DateTime                 │
│ + IsRevoked  : bool                     │
│ + CreatedAt  : DateTime                 │
└─────────────────────────────────────────┘

┌──────────────────────────────┐   ┌──────────────────────────────┐
│       <<interface>>          │   │       <<interface>>          │
│       IUserRepository        │   │       ITokenService          │
├──────────────────────────────┤   ├──────────────────────────────┤
│ + EmailExistsAsync()         │   │ + GenerateAccessToken()      │
│ + AddUserAsync()             │   │ + GenerateRefreshToken()     │
│ + GetByEmailAsync()          │   │ + ValidateAccessToken()      │
│ + GetByIdAsync()             │   └──────────────────────────────┘
│ + UpdateAsync()              │
└──────────────────────────────┘   ┌──────────────────────────────┐
                                   │       <<interface>>          │
┌──────────────────────────────┐   │   IRedisBlocklistService     │
│       <<interface>>          │   ├──────────────────────────────┤
│   IRefreshTokenRepository    │   │ + AddToBlocklistAsync()      │
├──────────────────────────────┤   │ + IsBlockedAsync()           │
│ + AddTokenAsync()            │   └──────────────────────────────┘
│ + GetActiveTokensAsync()     │
│ + UpdateAsync()              │   ┌──────────────────────────────┐
└──────────────────────────────┘   │       <<interface>>          │
                                   │   IEmailService (Identity)   │
                                   ├──────────────────────────────┤
                                   │ + SendVerificationEmailAsync()│
                                   │ + SendPasswordResetAsync()   │
                                   │ + SendWelcomeEmailAsync()    │
                                   └──────────────────────────────┘
""".strip().splitlines():
    mono(l)
gap()

# ── 3.2 JobCatalogService ─────────────────────────────────────────────────────
h("3.2  JobCatalogService", 2)
for l in """
┌──────────────────────────────────┐
│           <<entity>>             │
│             Company              │
├──────────────────────────────────┤
│ + Id          : Guid             │
│ + Name        : string           │
│ + Description : string?          │
│ + Website     : string?          │
│ + LogoUrl     : string?          │
│ + Industry    : string?          │
│ + Location    : string?          │
│ + CreatedAt   : DateTime         │
│ + Jobs        : ICollection<Job> │
│ + Reviews     : ICollection<>    │
└──────────────────────────────────┘
        1 │              1 │
          │ has many       │ has many
        * ▼              * ▼
┌──────────────────────┐  ┌──────────────────────────┐
│      <<entity>>      │  │        <<entity>>         │
│         Job          │  │      CompanyReview        │
├──────────────────────┤  ├──────────────────────────┤
│ + Id              Guid│  │ + Id          : Guid      │
│ + CompanyId       Guid│  │ + CompanyId   : Guid (FK) │
│ + PostedByRecruiter   │  │ + ReviewerId  : Guid      │
│ + Title       string  │  │ + Rating      : int       │
│ + Description text    │  │ + Comment     : string?   │
│ + Location    string  │  │ + IsApproved  : bool      │
│ + JobType     string  │  │ + CreatedAt   : DateTime  │
│ + SalaryMin   decimal?│  └──────────────────────────┘
│ + SalaryMax   decimal?│
│ + ModerationStatus    │
│ + CreatedAt   DateTime│
│ + UpdatedAt   DateTime│
└──────────────────────┘

┌──────────────────────────────┐   ┌──────────────────────────────┐
│       <<interface>>          │   │       <<interface>>          │
│       IJobRepository         │   │     ICompanyRepository       │
├──────────────────────────────┤   ├──────────────────────────────┤
│ + GetJobByIdAsync()          │   │ + GetCompanyByIdAsync()      │
│ + GetJobsByCompanyAsync()    │   │ + GetApprovedReviewsAsync()  │
│ + AddJobAsync()              │   │ + AddReviewAsync()           │
│ + UpdateJobAsync()           │   │ + AddCompanyAsync()          │
└──────────────────────────────┘   └──────────────────────────────┘

┌──────────────────────────────────────────────────────┐
│                   <<interface>>                      │
│               IElasticsearchService                  │
├──────────────────────────────────────────────────────┤
│ + SearchAsync(q, location, jobType, salary, page)    │
│ + IndexJobAsync(job, companyName)                    │
│ + UpdateJobAsync(job, companyName)                   │
│ + RemoveJobAsync(jobId)                              │
└──────────────────────────────────────────────────────┘
""".strip().splitlines():
    mono(l)
gap()

# ── 3.3 ApplicationService ────────────────────────────────────────────────────
h("3.3  ApplicationService", 2)
for l in """
┌──────────────────────────────────────┐
│              <<enum>>                │
│           ApplicationStatus         │
├──────────────────────────────────────┤
│  Submitted                           │
│  Reviewed                            │
│  Shortlisted                         │
│  Rejected                            │
│  Withdrawn                           │
└──────────────────────────────────────┘

┌──────────────────────────────────────┐
│             <<entity>>               │
│             Application              │
├──────────────────────────────────────┤
│ + Id              : Guid             │
│ + JobSeekerId     : Guid             │
│ + JobId           : Guid             │
│ + JobSeekerEmail  : string           │
│ + Status          : ApplicationStatus│
│ + CoverLetter     : string?          │
│ + AppliedAt       : DateTime         │
│ + UpdatedAt       : DateTime         │
└──────────────────────────────────────┘

┌──────────────────────────────────────┐
│             <<event>>                │
│      ApplicationStatusChangedEvent   │
├──────────────────────────────────────┤
│ + ApplicationId : Guid               │
│ + JobSeekerId   : Guid               │
│ + JobId         : Guid               │
│ + OldStatus     : string             │
│ + NewStatus     : string             │
│ + ChangedAt     : DateTime           │
└──────────────────────────────────────┘

┌──────────────────────────────┐   ┌──────────────────────────────┐
│       <<interface>>          │   │       <<interface>>          │
│    IApplicationRepository    │   │   IApplicationStateMachine   │
├──────────────────────────────┤   ├──────────────────────────────┤
│ + HasAppliedAsync()          │   │ + ValidateTransition()       │
│ + AddApplicationAsync()      │   │ + GetAllowedTransitions()    │
│ + GetApplicationByIdAsync()  │   └──────────────────────────────┘
│ + GetApplicationsBySeeker()  │
│ + GetApplicationsByJob()     │   ┌──────────────────────────────┐
│ + UpdateApplicationAsync()   │   │       <<interface>>          │
└──────────────────────────────┘   │   IEmailService (App)        │
                                   ├──────────────────────────────┤
┌──────────────────────────────┐   │ + SendApplicationSubmitted() │
│       <<interface>>          │   │ + SendStatusChanged()        │
│      IRabbitMqPublisher      │   │ + SendApplicationWithdrawn() │
├──────────────────────────────┤   └──────────────────────────────┘
│ + PublishStatusChangedAsync()│
└──────────────────────────────┘
""".strip().splitlines():
    mono(l)
gap()

# ── 3.4 ResumeService ─────────────────────────────────────────────────────────
h("3.4  ResumeService", 2)
for l in """
┌──────────────────────────────────────────┐
│                <<entity>>                │
│                  Resume                  │
├──────────────────────────────────────────┤
│ + Id             : Guid                  │
│ + OwnerId        : Guid                  │
│ + Title          : string                │
│ + Summary        : string?               │
│ + TemplateId     : string                │
│ + Certifications : string?               │
│ + CreatedAt      : DateTime              │
│ + UpdatedAt      : DateTime              │
│ + Educations     : ICollection<>         │
│ + Experiences    : ICollection<>         │
│ + Skills         : ICollection<>         │
│ + Projects       : ICollection<>         │
└──────────────────────────────────────────┘
   1│      1│      1│      1│
    │       │       │       │  (all CASCADE DELETE)
   *▼      *▼      *▼      *▼
┌──────────┐ ┌──────────┐ ┌──────────┐ ┌──────────┐
│Education │ │Experience│ │  Skill   │ │ Project  │
├──────────┤ ├──────────┤ ├──────────┤ ├──────────┤
│Id   Guid │ │Id   Guid │ │Id   Guid │ │Id   Guid │
│ResumeId  │ │ResumeId  │ │ResumeId  │ │ResumeId  │
│Institution│ │Company   │ │Name      │ │Name      │
│Degree    │ │JobTitle  │ │Level?    │ │Description│
│FieldOfStudy│ │Description│ └──────────┘ │Url?      │
│StartDate │ │StartDate │              └──────────┘
│EndDate?  │ │EndDate?  │
└──────────┘ │IsCurrentRole│
             └──────────┘

┌──────────────────────────────┐   ┌──────────────────────────────┐
│          <<entity>>          │   │          <<entity>>          │
│     ResumeUnlockRequest      │   │        UnlockedResume        │
├──────────────────────────────┤   ├──────────────────────────────┤
│ + Id          : Guid         │   │ + Id          : Guid         │
│ + RecruiterId : Guid         │   │ + RecruiterId : Guid         │
│ + ResumeId    : Guid         │   │ + ResumeId    : Guid         │
│ + Status      : string       │   │ + UnlockedAt  : DateTime     │
│ + CreatedAt   : DateTime     │   └──────────────────────────────┘
│ + UpdatedAt   : DateTime     │
└──────────────────────────────┘

┌──────────────────────────────┐   ┌──────────────────────────────┐
│       <<interface>>          │   │       <<interface>>          │
│      IResumeRepository       │   │     IPdfGeneratorService     │
├──────────────────────────────┤   ├──────────────────────────────┤
│ + GetResumesByOwnerAsync()   │   │ + GeneratePdf(resume)        │
│ + GetResumeWithDetailsAsync()│   └──────────────────────────────┘
│ + AddResumeAsync()           │
│ + RemoveDetailsAsync()       │
│ + SaveChangesAsync()         │
└──────────────────────────────┘
""".strip().splitlines():
    mono(l)
gap()

# ── 3.5 PaymentService ────────────────────────────────────────────────────────
h("3.5  PaymentService", 2)
for l in """
┌──────────────────────────────────────┐
│              <<enum>>                │
│           TransactionType            │
├──────────────────────────────────────┤
│  Credit                              │
│  Debit                               │
└──────────────────────────────────────┘

┌──────────────────────────────────────┐
│             <<entity>>               │
│           RecruiterWallet            │
├──────────────────────────────────────┤
│ + Id             : Guid              │
│ + RecruiterId    : Guid  (UNIQUE)    │
│ + PointsBalance  : int               │
│ + CreatedAt      : DateTime          │
│ + UpdatedAt      : DateTime          │
│ + RowVersion     : byte[]  (concurrency token)
│ + Transactions   : ICollection<>     │
└──────────────────────────────────────┘
              1 │
                │ has many
              * ▼
┌──────────────────────────────────────┐
│             <<entity>>               │
│             Transaction              │
├──────────────────────────────────────┤
│ + Id               : Guid            │
│ + WalletId         : Guid  (FK)      │
│ + Type             : TransactionType │
│ + Points           : int             │
│ + Description      : string?         │
│ + IdempotencyKey   : string? (UNIQUE)│
│ + RazorpayPaymentId: string?         │
│ + AmountPaid       : decimal?        │
│ + Currency         : string?         │
│ + CreatedAt        : DateTime        │
└──────────────────────────────────────┘

┌──────────────────────────────────────┐
│             <<entity>>               │
│         PointsDeductionRule          │
├──────────────────────────────────────┤
│ + Id       : Guid                    │
│ + Action   : string  (ResumeView|ContactUnlock)
│ + Points   : int                     │
│ + IsActive : bool                    │
└──────────────────────────────────────┘

┌──────────────────────────────────────┐
│             <<saga state>>           │
│        UnlockResumeSagaState         │
├──────────────────────────────────────┤
│ + CorrelationId  : Guid  (PK)        │
│ + CurrentState   : string            │
│ + RecruiterId    : Guid              │
│ + ResumeId       : Guid              │
│ + PointsToDeduct : int               │
│ + CreatedAt      : DateTime          │
│ + UpdatedAt      : DateTime?         │
└──────────────────────────────────────┘

┌──────────────────────────────┐   ┌──────────────────────────────┐
│       <<interface>>          │   │       <<interface>>          │
│      IPaymentRepository      │   │        IWalletService        │
├──────────────────────────────┤   ├──────────────────────────────┤
│ + GetWalletAsync()           │   │ + GetBalanceAsync()          │
│ + EnsureWalletExistsAsync()  │   │ + DeductAsync()              │
│ + UpdateWalletAsync()        │   │ + CreditAsync()              │
│ + AddTransactionAsync()      │   │ + EnsureWalletExistsAsync()  │
│ + GetDeductionRuleAsync()    │   └──────────────────────────────┘
│ + GetTransactionsAsync()     │
│ + GetRevenueReportAsync()    │   ┌──────────────────────────────┐
└──────────────────────────────┘   │       <<interface>>          │
                                   │      IRazorpayService        │
                                   ├──────────────────────────────┤
                                   │ + CreateOrderAsync()         │
                                   │ + VerifyWebhookSignature()   │
                                   │ + VerifyPaymentSignature()   │
                                   └──────────────────────────────┘
""".strip().splitlines():
    mono(l)
gap()

# ── 3.6 AdminService ──────────────────────────────────────────────────────────
h("3.6  AdminService", 2)
for l in """
┌──────────────────────────────────────┐
│             <<entity>>               │
│              AuditLog                │
├──────────────────────────────────────┤
│ + Id         : Guid                  │
│ + AdminId    : Guid                  │
│ + Action     : string                │
│ + TargetType : string                │
│ + TargetId   : Guid                  │
│ + Details    : string?               │
│ + CreatedAt  : DateTime              │
└──────────────────────────────────────┘

<<read models — no DB writes, cross-DB reads>>
┌──────────────────────────────┐   ┌──────────────────────────────┐
│       <<interface>>          │   │       <<interface>>          │
│      IUserRepository         │   │       IJobRepository         │
│       (AdminService)         │   │       (AdminService)         │
├──────────────────────────────┤   ├──────────────────────────────┤
│ + GetUsersAsync()            │   │ + GetModerationQueueAsync()  │
│ + GetUserByIdAsync()         │   │ + GetJobByIdAsync()          │
│ + UpdateUserAsync()          │   │ + UpdateJobAsync()           │
└──────────────────────────────┘   └──────────────────────────────┘

┌──────────────────────────────┐   ┌──────────────────────────────┐
│       <<interface>>          │   │       <<interface>>          │
│     IAuditLogRepository      │   │     IPaymentRepository       │
│       (AdminService)         │   │       (AdminService)         │
├──────────────────────────────┤   ├──────────────────────────────┤
│ + AddAuditLogAsync()         │   │ + GetTransactionsAsync()     │
│ + GetAuditLogsAsync()        │   │ + GetRevenueAsync()          │
└──────────────────────────────┘   └──────────────────────────────┘
""".strip().splitlines():
    mono(l)
gap()

# ══════════════════════════════════════════════════════════════════════════════
# 4. UML COMPONENT DIAGRAM
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h("4. UML Component Diagram")
para("Shows how all components and services connect at the system level.", italic=True, size=10)
gap()
for l in """
 ┌─────────────────────────────────────────────────────────────────────────┐
 │                          Client Applications                            │
 │                    (Web Browser / Mobile / Postman)                     │
 └───────────────────────────────┬─────────────────────────────────────────┘
                                 │ HTTP
                                 ▼
 ┌───────────────────────────────────────────────────────────────────────┐
 │                          ApiGateway  :5000                            │
 │   [Ocelot Router]  [JWT RS256 Validator]  [Rate Limiter 100 req/min]  │
 └──┬──────────┬──────────┬──────────┬──────────┬──────────┬────────────┘
    │          │          │          │          │          │
    ▼          ▼          ▼          ▼          ▼          ▼
 ┌──────┐  ┌──────┐  ┌──────┐  ┌──────┐  ┌──────┐  ┌──────┐
 │:5001 │  │:5002 │  │:5003 │  │:5004 │  │:5005 │  │:5006 │
 │Ident-│  │JobCat│  │Appli-│  │Resum-│  │Payme-│  │Admin │
 │ity   │  │alog  │  │cation│  │e     │  │nt    │  │      │
 └──┬───┘  └──┬───┘  └──┬───┘  └──┬───┘  └──┬───┘  └──┬───┘
    │         │          │         │          │          │
    ▼         ▼          ▼         ▼          ▼          ▼
 ┌──────┐  ┌──────┐  ┌──────┐  ┌──────┐  ┌──────┐  ┌──────┐
 │Ident-│  │JobDb │  │Appli-│  │Resum-│  │Payme-│  │Admin-│
 │ityDb │  │      │  │ation-│  │eDb   │  │ntDb  │  │Db    │
 └──┬───┘  └──┬───┘  │Db    │  └──────┘  └──────┘  └──────┘
    │         │      └──────┘
    ▼         ▼
 ┌──────┐  ┌──────────────┐
 │Redis │  │Elasticsearch │
 │(JTI  │  │(Job Index)   │
 │block)│  └──────────────┘
 └──────┘

 Cross-service communication:
 ┌─────────────────┐  HTTP POST /deduct  ┌─────────────────┐
 │  ResumeService  │ ──────────────────► │  PaymentService │
 └─────────────────┘                     └─────────────────┘

 ┌─────────────────┐  RabbitMQ fanout    ┌─────────────────┐
 │ApplicationService│ ─────────────────► │  (subscribers)  │
 │  [Publisher]    │  application-events │  Email / Push   │
 └─────────────────┘                     └─────────────────┘

 ┌─────────────────┐  Direct DB reads    ┌─────────────────┐
 │  AdminService   │ ──────────────────► │IdentityDb/JobDb │
 │                 │                     │  /PaymentDb     │
 └─────────────────┘                     └─────────────────┘
""".strip().splitlines():
    mono(l)
gap()

# ══════════════════════════════════════════════════════════════════════════════
# 5. UML USE CASE DIAGRAM
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h("5. UML Use Case Diagram")
para("Actors and their interactions with the system.", italic=True, size=10)
gap()
for l in """
                    ┌─────────────────────────────────────────────────────┐
                    │                   Job Portal System                 │
                    │                                                     │
  ┌──────────┐      │  ┌─────────────────────────────────────────────┐   │
  │          │      │  │  <<use case>>  Register / Login / Logout    │   │
  │JobSeeker │──────┼─►│  <<use case>>  Verify Email                 │   │
  │          │      │  │  <<use case>>  Search Jobs                  │   │
  └──────────┘      │  │  <<use case>>  Apply for Job                │   │
       │            │  │  <<use case>>  Withdraw Application         │   │
       │            │  │  <<use case>>  Track Application Status     │   │
       │            │  │  <<use case>>  Create / Update Resume       │   │
       │            │  │  <<use case>>  Download Resume as PDF       │   │
       │            │  └─────────────────────────────────────────────┘   │
       │            │                                                     │
  ┌──────────┐      │  ┌─────────────────────────────────────────────┐   │
  │          │      │  │  <<use case>>  Register / Login             │   │
  │Recruiter │──────┼─►│  <<use case>>  Post / Edit Job              │   │
  │          │      │  │  <<use case>>  View Applicants              │   │
  └──────────┘      │  │  <<use case>>  Update Application Status    │   │
       │            │  │  <<use case>>  Shortlist Candidate          │   │
       │            │  │  <<use case>>  Purchase Points (Razorpay)   │   │
       │            │  │  <<use case>>  View Resume (costs points)   │   │
       │            │  │  <<use case>>  Check Wallet Balance         │   │
       │            │  └─────────────────────────────────────────────┘   │
       │            │                                                     │
  ┌──────────┐      │  ┌─────────────────────────────────────────────┐   │
  │          │      │  │  <<use case>>  Approve / Flag Jobs          │   │
  │  Admin   │──────┼─►│  <<use case>>  Manage Users (role/delete)   │   │
  │          │      │  │  <<use case>>  View Audit Logs              │   │
  └──────────┘      │  │  <<use case>>  View Revenue Reports         │   │
                    │  └─────────────────────────────────────────────┘   │
                    │                                                     │
  ┌──────────┐      │  ┌─────────────────────────────────────────────┐   │
  │          │      │  │  <<use case>>  Send Webhook (payment done)  │   │
  │Razorpay  │──────┼─►│  <<use case>>  Credit Wallet (idempotent)   │   │
  │          │      │  └─────────────────────────────────────────────┘   │
  └──────────┘      │                                                     │
                    └─────────────────────────────────────────────────────┘
""".strip().splitlines():
    mono(l)
gap()

# ══════════════════════════════════════════════════════════════════════════════
# 6. ACTIVITY DIAGRAMS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h("6. Activity Diagrams")

h("6.1  User Registration & Email Verification", 2)
for l in """
● START
│
▼
[User submits POST /auth/register  { email, password, role }]
│
▼
<Validate input (FluentValidation)>
├─ INVALID ──► Return 400 Bad Request
▼  VALID
[Hash password with BCrypt (work factor 10)]
│
▼
[Save User to IdentityDb  (EmailVerified = false)]
│
▼
[Generate EmailVerificationToken (GUID), set expiry]
│
▼
[Send verification email via SMTP]
│
▼
[User clicks link → GET /auth/verify-email?token=...]
│
▼
<Token found & not expired?>
├─ NO  ──► Return 400 Invalid Token
▼  YES
[Set EmailVerified = true, clear token]
│
▼
Return 200 OK
│
● END
""".strip().splitlines():
    mono(l)
gap()

h("6.2  Login & JWT Issuance", 2)
for l in """
● START
│
▼
[User submits POST /auth/login  { email, password }]
│
▼
<User exists?>
├─ NO  ──► Return 401
▼  YES
<Account locked (LockoutEnd > now)?>
├─ YES ──► Return 423 Locked (remaining seconds)
▼  NO
<Password matches BCrypt hash?>
├─ NO  ──► Increment FailedLoginAttempts
│          <Attempts >= 5?> ──► Lock account 15 min
│          Return 401
▼  YES
[Reset FailedLoginAttempts = 0]
│
▼
<EmailVerified = true?>
├─ NO  ──► Return 403 Email not verified
▼  YES
[Generate JWT (RS256, 15 min) with sub, email, role, jti]
│
▼
[Generate RefreshToken, hash with BCrypt, store in DB (7 day expiry)]
│
▼
Return 200 OK  { accessToken, refreshToken }
│
● END
""".strip().splitlines():
    mono(l)
gap()

h("6.3  Job Posting & Admin Moderation", 2)
for l in """
● START
│
▼
[Recruiter submits POST /jobs  (Role = Recruiter)]
│
▼
<Company exists?>
├─ NO  ──► Return 400
▼  YES
[Create Job  (ModerationStatus = "Pending")]
│
▼
Return 201 Created  { jobId }
│
▼
[Admin opens moderation queue]
│
▼
<Admin decision>
├─ APPROVE ──► [Set ModerationStatus = "Approved"]
│              [Index job in Elasticsearch]
│              [Write AuditLog]
│              Return 200 OK
│
└─ FLAG    ──► [Set ModerationStatus = "Flagged"]
               [Remove job from Elasticsearch]
               [Write AuditLog]
               Return 200 OK
│
● END
""".strip().splitlines():
    mono(l)
gap()

h("6.4  Job Application Lifecycle", 2)
for l in """
● START
│
▼
[JobSeeker submits POST /applications  { jobId, coverLetter, email }]
│
▼
<Already applied?>
├─ YES ──► Return 409 Conflict
▼  NO
[Create Application  (Status = "Submitted")]
│
▼
[Publish ApplicationStatusChanged → RabbitMQ]
│
▼
[Send "Application Submitted" email to JobSeeker]
│
▼
Return 201 Created  { applicationId }
│
▼
[Recruiter: PATCH /applications/{id}/status  { newStatus }]
│
▼
<Valid state transition?>
├─ NO  ──► Return 422
▼  YES
[Update Status in DB]
│
▼
[Publish ApplicationStatusChanged → RabbitMQ]
│
▼
[Send status-change email to JobSeeker]
│
▼
<JobSeeker withdraws? DELETE /applications/{id}>
├─ Status != Submitted ──► Return 422
▼  Status = Submitted
[Set Status = "Withdrawn"]
│
▼
[Send "Application Withdrawn" email]
│
▼
Return 200 OK
│
● END
""".strip().splitlines():
    mono(l)
gap()

h("6.5  Resume Create & Update", 2)
for l in """
● START (Create)
│
▼
[JobSeeker submits POST /resumes  { title, summary, templateId, sections }]
│
▼
[Create Resume + child rows in single INSERT batch]
│
▼
Return 201 Created  { resumeId }
│
● END (Create)

● START (Update)
│
▼
[JobSeeker submits PUT /resumes/{id}]
│
▼
<Resume exists & owned by caller?>
├─ NO  ──► Return 404 / 403
▼  YES
[ExecuteDeleteAsync on all child tables WHERE ResumeId = id  (direct SQL)]
│
▼
[ExecuteUpdateAsync on Resumes  (title, summary, templateId, updatedAt)]
│
▼
[AddRange new child rows]
│
▼
[SaveChangesAsync  (INSERT only, no concurrency conflict)]
│
▼
Return 200 OK  { resumeId }
│
● END (Update)
""".strip().splitlines():
    mono(l)
gap()

h("6.6  Recruiter Wallet & Points Deduction", 2)
for l in """
● START
│
▼
[Recruiter: POST /payments/wallet/purchase  { amountInPaise }]
│
▼
[Create Razorpay order]
│
▼
Return 200 OK  { orderId }  ← frontend opens Razorpay checkout
│
▼
[Razorpay sends webhook POST /payments/wallet/webhook]
│
▼
<HMAC-SHA256 signature valid?>
├─ NO  ──► Return 400
▼  YES
<IdempotencyKey already processed?>
├─ YES ──► Return 200 (no-op)
▼  NO
[Credit wallet: PointsBalance += amount]
[Insert Transaction with IdempotencyKey]
│
▼
Return 200 OK
│
▼
[Recruiter: GET /resumes/{id}/view]
│
▼
[ResumeService → POST /payments/wallet/deduct  { action:"ResumeView" }]
│
▼
<PointsBalance >= 10?>
├─ NO  ──► Return 402 Payment Required
▼  YES
[Deduct 10 points  (RowVersion optimistic concurrency, retry x3)]
│
▼
Return 200 OK → ResumeService returns full resume
│
● END
""".strip().splitlines():
    mono(l)
gap()

# ══════════════════════════════════════════════════════════════════════════════
# 7. SEQUENCE DIAGRAMS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h("7. Sequence Diagrams")

h("7.1  User Registration & Email Verification", 2)
for l in """
Client        ApiGateway      IdentityService     IdentityDb       SMTP
  │                │                 │                  │             │
  │─POST /register──────────────────►│                  │             │
  │                │─route──────────►│                  │             │
  │                │                 │─validate          │             │
  │                │                 │─bcrypt hash       │             │
  │                │                 │─INSERT User──────►│             │
  │                │                 │◄──────────────────│             │
  │                │                 │─send verify email──────────────►│
  │◄───────────────────────────────201 { userId }        │             │
  │                │                 │                  │             │
  │─GET /verify-email?token=─────────►                  │             │
  │                │─route──────────►│                  │             │
  │                │                 │─UPDATE User (EmailVerified=true)►│
  │◄───────────────────────────────200 OK               │             │
""".strip().splitlines():
    mono(l)
gap()

h("7.2  Login & Logout", 2)
for l in """
Client        ApiGateway      IdentityService     IdentityDb       Redis
  │                │                 │                  │             │
  │─POST /login─────────────────────►│                  │             │
  │                │─route──────────►│                  │             │
  │                │                 │─SELECT User──────►│             │
  │                │                 │◄──────────────────│             │
  │                │                 │─verify BCrypt     │             │
  │                │                 │─sign JWT (RS256)  │             │
  │                │                 │─INSERT RefreshToken────────────►│
  │◄───────────────────────────────200 { accessToken, refreshToken }   │
  │                │                 │                  │             │
  │─POST /logout (Bearer JWT)────────►                  │             │
  │                │─route──────────►│                  │             │
  │                │                 │─SET jti:{jti} EX {ttl}─────────►│
  │◄───────────────────────────────200 OK               │             │
""".strip().splitlines():
    mono(l)
gap()

h("7.3  Job Search via Elasticsearch", 2)
for l in """
Client       ApiGateway    JobCatalogService    Elasticsearch
  │               │                │                  │
  │─GET /jobs/search?q=dev─────────►                  │
  │               │─route─────────►│                  │
  │               │                │─SearchAsync(q, filters, page)────►│
  │               │                │◄──────────────────────────────────│
  │◄──────────────────────────────200 { total, jobs[] }                │
""".strip().splitlines():
    mono(l)
gap()

h("7.4  Post Job & Admin Approval", 2)
for l in """
Recruiter   ApiGateway  JobCatalogService   JobDb      Elasticsearch
  │              │               │              │              │
  │─POST /jobs───────────────────►             │              │
  │              │─route────────►│             │              │
  │              │               │─INSERT Job (Pending)──────►│
  │◄─────────────────────────────201 { jobId } │              │
  │              │               │             │              │
Admin       ApiGateway   AdminService       JobDb      Elasticsearch
  │              │               │              │              │
  │─POST /admin/jobs/{id}/approve──────────────►              │
  │              │─route────────►│             │              │
  │              │               │─UPDATE Job (Approved)─────►│
  │              │               │─INSERT AuditLog───────────►│
  │              │               │─IndexJobAsync──────────────────────►│
  │◄─────────────────────────────200 OK        │              │◄───────│
""".strip().splitlines():
    mono(l)
gap()

h("7.5  Job Application & Email Notification", 2)
for l in """
JobSeeker  ApiGateway  ApplicationService  ApplicationDb  RabbitMQ  EmailService
  │             │               │                │             │           │
  │─POST /applications──────────►              │             │           │
  │             │─route────────►│              │             │           │
  │             │               │─HasApplied?─►│             │           │
  │             │               │◄─────────────│             │           │
  │             │               │─INSERT Application────────►│           │
  │             │               │─PublishStatusChanged───────────────────►│
  │             │               │─SendApplicationSubmitted───────────────►│
  │◄────────────────────────────201 { applicationId }        │           │
  │             │               │                │             │           │
  │             │               │  [Recruiter shortlists]     │           │
  │             │               │─UPDATE Status (Shortlisted)►│           │
  │             │               │─PublishStatusChanged───────────────────►│
  │             │               │─SendStatusChanged (Shortlisted)────────►│
  │◄────────────────────────────200 OK          │             │           │
""".strip().splitlines():
    mono(l)
gap()

h("7.6  Resume View with Points Deduction", 2)
for l in """
Recruiter  ApiGateway  ResumeService   PaymentService   PaymentDb   ResumeDb
  │             │              │               │              │          │
  │─GET /resumes/{id}/view──────►             │              │          │
  │             │─route───────►│              │              │          │
  │             │              │─POST /wallet/deduct─────────►          │
  │             │              │              │─SELECT Wallet────────── ►│
  │             │              │              │◄──────────────────────── │
  │             │              │              │─balance < 10?            │
  │             │              │◄──402────────│              │          │
  │◄────────────────────────────402 Payment Required         │          │
  │             │              │              │  balance >= 10           │
  │             │              │              │─UPDATE Wallet (balance-10)►│
  │             │              │              │─INSERT Transaction───────►│
  │             │              │◄──200────────│              │          │
  │             │              │─SELECT Resume with sections─────────────►│
  │◄────────────────────────────200 { full resume data }     │          │
""".strip().splitlines():
    mono(l)
gap()

h("7.7  Razorpay Webhook & Wallet Credit", 2)
for l in """
Razorpay     PaymentService       PaymentDb
  │                │                   │
  │─POST /wallet/webhook──────────────►│
  │                │─verify HMAC-SHA256 │
  │                │  INVALID ──► 400  │
  │                │─SELECT Transaction (IdempotencyKey)──────────────►│
  │                │◄──────────────────│                               │
  │                │  EXISTS ──► 200 (idempotent no-op)                │
  │                │─SELECT Wallet─────────────────────────────────────►│
  │                │◄──────────────────│                               │
  │                │─UPDATE Wallet (balance += points)─────────────────►│
  │                │─INSERT Transaction (IdempotencyKey)────────────────►│
  │◄──────────────200 OK              │                               │
""".strip().splitlines():
    mono(l)
gap()

h("7.8  Token Refresh", 2)
for l in """
Client       ApiGateway      IdentityService       IdentityDb
  │               │                  │                   │
  │─POST /auth/refresh-token─────────►                   │
  │               │─route───────────►│                   │
  │               │                  │─SELECT RefreshToken (by hash)──►│
  │               │                  │◄──────────────────│              │
  │               │                  │─validate not expired, not revoked│
  │               │                  │─UPDATE old token (IsRevoked=true)►│
  │               │                  │─sign new JWT (RS256)             │
  │               │                  │─INSERT new RefreshToken─────────►│
  │◄──────────────────────────────200 { accessToken, refreshToken }     │
""".strip().splitlines():
    mono(l)
gap()

# ══════════════════════════════════════════════════════════════════════════════
# 8. APPLICATION STATE MACHINE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h("8. Application Status State Machine")
gap()
for l in """
  ┌──────────────────────────────────────────────────────────────────────┐
  │                                                                      │
  │   ●──► Submitted ─────────────────────────────────────► Withdrawn   │
  │            │                                           (JobSeeker)  │
  │            │                                                         │
  │            ├──────────────────► Reviewed                             │
  │            │                       │                                 │
  │            │                       ├──────────────► Shortlisted      │
  │            │                       │                    │            │
  │            │                       ▼                    ▼            │
  │            └───────────────────────┴──────────────► Rejected ●      │
  │                                                                      │
  └──────────────────────────────────────────────────────────────────────┘
""".strip().splitlines():
    mono(l)
gap()
table([
    ("Current Status","Allowed Next Statuses","Who Can Trigger"),
    ("Submitted","Reviewed, Rejected, Withdrawn","Recruiter (Reviewed/Rejected), JobSeeker (Withdrawn)"),
    ("Reviewed","Shortlisted, Rejected","Recruiter"),
    ("Shortlisted","Rejected","Recruiter"),
    ("Rejected","— (terminal)","—"),
    ("Withdrawn","— (terminal)","—"),
])
gap()

# ══════════════════════════════════════════════════════════════════════════════
# 9. EMAIL NOTIFICATION MATRIX
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h("9. Email Notification Matrix")
gap()
table([
    ("Trigger","Subject","Recipient","Triggered By"),
    ("POST /applications","Application Submitted Successfully","JobSeeker","JobSeeker"),
    ("PATCH .../status → Reviewed","Application Status Updated: Reviewed","JobSeeker","Recruiter"),
    ("PATCH .../status → Shortlisted","Application Status Updated: Shortlisted","JobSeeker","Recruiter"),
    ("PATCH .../status → Rejected","Application Status Updated: Rejected","JobSeeker","Recruiter"),
    ("PATCH .../shortlist","Application Status Updated: Shortlisted","JobSeeker","Recruiter"),
    ("DELETE /applications/{id}","Application Withdrawn","JobSeeker","JobSeeker"),
])
gap()
para("SMTP: smtp.gmail.com : 587 (STARTTLS) — configured in ApplicationService/appsettings.json", size=10)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
doc.save("JobPortalSystem_Design_Document.docx")
print("Done → JobPortalSystem_Design_Document.docx")
